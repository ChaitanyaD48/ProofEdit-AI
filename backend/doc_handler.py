import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from models import FinalDocumentData

def extract_text_from_docx(docx_file_stream: io.BytesIO) -> str:
    """Reads a .docx file and returns its raw text."""
    try:
        document = Document(docx_file_stream)
        return '\n'.join([para.text for para in document.paragraphs])
    except Exception as e:
        raise ValueError(f"Could not read the docx file: {e}")

def create_final_docx(params: FinalDocumentData) -> io.BytesIO:
    """Creates the final, formatted .docx file from the edited text and glossary."""
    document = Document()
    
    # --- 1. Apply Global Formatting ---
    style = document.styles['Normal']
    font = style.font
    font.name = params.font_family
    font.size = Pt(params.font_size)
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = params.line_spacing

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(params.margin_top)
        section.bottom_margin = Inches(params.margin_bottom)
        section.left_margin = Inches(params.margin_left)
        section.right_margin = Inches(params.margin_right)

    # --- 2. Parse and Add Manuscript Content ---
    content_without_markers = params.edited_manuscript
    pattern = re.compile(r'(\[(H1|H2|SHLOKA|TRANSLATION)\].*?\[/\2\])', re.DOTALL)
    
    cursor = 0
    for match in pattern.finditer(content_without_markers):
        plain_text = content_without_markers[cursor:match.start()]
        if plain_text.strip():
            document.add_paragraph(plain_text.strip())

        tag = match.group(2)
        content = match.group(1)[len(tag)+2:-len(tag)-3].strip()

        if tag == 'H1':
            p = document.add_paragraph()
            run = p.add_run(content)
            run.font.size = Pt(params.heading1.font_size)
            run.bold = params.heading1.bold
        elif tag == 'H2':
            p = document.add_paragraph()
            run = p.add_run(content)
            run.font.size = Pt(params.heading2.font_size)
            run.bold = params.heading2.bold
        elif tag == 'SHLOKA':
            p = document.add_paragraph()
            lines = content.split('\n')
            for i, line in enumerate(lines):
                if i > 0:
                    p.add_run().add_break()
                run = p.add_run(line.strip())
                run.italic = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif tag == 'TRANSLATION':
            p = document.add_paragraph(content)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        cursor = match.end()

    remaining_text = content_without_markers[cursor:]
    if remaining_text.strip():
        cleaned_remaining = re.sub(r'\[/?(ITALIC|CITE:.*?)\]', '', remaining_text)
        document.add_paragraph(cleaned_remaining.strip())

    # --- 3. Add Glossary Section ---
    if params.glossary:
        document.add_page_break()
        document.add_heading('Glossary', level=1)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text = 'Term', 'Transliteration', 'Translation', 'Context/Citation'
        for item in params.glossary:
            row_cells = table.add_row().cells
            row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text = item.term, item.transliteration, item.translation, item.context or ''

    # --- 4. Save to Memory Buffer ---
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

